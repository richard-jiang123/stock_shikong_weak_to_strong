#!/usr/bin/env python3
"""Extract text and images from all docx files in shikong_fufei."""
import os
import base64
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

DOC_DIR = "/home/jzc/wechat_text/shikong_fufei"
OUT_DIR = "/home/jzc/wechat_text/shikong_fufei/extracted"
os.makedirs(OUT_DIR, exist_ok=True)

for fname in sorted(os.listdir(DOC_DIR)):
    if not fname.endswith('.docx'):
        continue
    fpath = os.path.join(DOC_DIR, fname)
    doc = Document(fpath)
    name = fname.replace('.docx', '')
    out_txt = os.path.join(OUT_DIR, f"{name}_text.txt")
    out_imgs = os.path.join(OUT_DIR, f"{name}_images")
    os.makedirs(out_imgs, exist_ok=True)

    # Extract text
    with open(out_txt, 'w', encoding='utf-8') as f:
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                f.write(f"[P{i}] {para.text.strip()}\n")
        # Also extract table text
        for ti, table in enumerate(doc.tables):
            f.write(f"\n--- TABLE {ti} ---\n")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                f.write(" | ".join(cells) + "\n")

    # Extract images
    img_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            ext = rel.target_ref.split('.')[-1]
            if ext not in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'tif', 'tiff', 'webp'):
                ext = 'png'
            img_data = rel.target_part.blob
            img_path = os.path.join(out_imgs, f"image_{img_count:03d}.{ext}")
            with open(img_path, 'wb') as f:
                f.write(img_data)
            print(f"  Image saved: {img_path} ({len(img_data)} bytes)")
            img_count += 1

    print(f"\n{name}: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {img_count} images")
    print(f"  Text: {out_txt}")
    print(f"  Images: {out_imgs}")
